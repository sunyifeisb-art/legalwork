from __future__ import annotations

import hashlib
import json
import math
import os
import re
import shutil
import uuid
from collections import Counter
from dataclasses import asdict, dataclass, field
from datetime import date, datetime, timezone
from pathlib import Path
from typing import Any

from redaction.pipeline import RedactionPipeline
from redaction.renderer import RedactionRenderer
from redaction.restorer import RedactionRestorer

try:
    from document.semantic.chunker import SemanticChunker
except Exception:  # pragma: no cover - optional in degraded installs
    SemanticChunker = None


ROOT_DIR = Path(__file__).resolve().parent.parent
GLOBAL_MATTER_ID = "__global__"
KNOWLEDGE_FILE_EXTENSIONS = {".md", ".txt", ".json", ".docx", ".pdf"}
MAX_KNOWLEDGE_FILE_BYTES = 2 * 1024 * 1024


def default_knowledge_roots() -> list[Path]:
    roots = [
        ROOT_DIR / "knowledge-base",
        ROOT_DIR / "projects" / "compliance" / "projects" / "data-compliance-ai-project-kit" / "knowledge-base",
        ROOT_DIR / "projects" / "compliance" / "skills" / "data-compliance-reviewer" / "references" / "knowledge-base",
    ]
    extra = os.environ.get("LEGALWORK_KNOWLEDGE_ROOTS", "")
    for item in extra.split(os.pathsep):
        if item.strip():
            roots.append(Path(item.strip()))
    return roots


def utc_now() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")


def stable_id(prefix: str) -> str:
    return f"{prefix}_{uuid.uuid4().hex[:12]}"


def sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8", errors="ignore")).hexdigest()


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as fh:
        for block in iter(lambda: fh.read(1024 * 1024), b""):
            h.update(block)
    return h.hexdigest()


def slugify_name(value: str, fallback: str = "matter") -> str:
    value = re.sub(r"[^\w\u4e00-\u9fff.-]+", "_", (value or "").strip(), flags=re.UNICODE)
    value = value.strip("._")
    return value or fallback


def load_json(path: Path, default: Any) -> Any:
    if not path.exists():
        return default
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return default


def write_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(path.suffix + ".tmp")
    tmp.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    tmp.replace(path)


def extract_text_from_path(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".csv", ".json", ".yaml", ".yml"}:
        for enc in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
            try:
                return path.read_text(encoding=enc)
            except UnicodeDecodeError:
                continue
        return path.read_text(encoding="utf-8", errors="replace")

    if suffix == ".docx":
        from docx import Document

        doc = Document(str(path))
        parts: list[str] = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    if suffix == ".pdf":
        try:
            import fitz

            with fitz.open(str(path)) as doc:
                return "\n".join(page.get_text("text") for page in doc)
        except Exception:
            return ""

    try:
        from skill_engine.intake import extract_text

        result = extract_text(path)
        return result.get("text", "") if result.get("success") else ""
    except Exception:
        return ""


def tokenize(text: str) -> list[str]:
    tokens = [t.lower() for t in re.findall(r"[A-Za-z0-9_]+|[\u4e00-\u9fff]", text or "")]
    cjk = re.findall(r"[\u4e00-\u9fff]{2,}", text or "")
    for block in cjk:
        tokens.extend(block[i : i + 2] for i in range(max(0, len(block) - 1)))
    return tokens


def vectorize(text: str) -> dict[str, float]:
    counts = Counter(tokenize(text))
    if not counts:
        return {}
    norm = math.sqrt(sum(v * v for v in counts.values())) or 1.0
    return {k: v / norm for k, v in counts.items()}


def cosine(a: dict[str, float], b: dict[str, float]) -> float:
    if not a or not b:
        return 0.0
    if len(a) > len(b):
        a, b = b, a
    return sum(v * b.get(k, 0.0) for k, v in a.items())


@dataclass
class Matter:
    matter_id: str
    name: str
    team_id: str = "default"
    source_channel: str = "local"
    created_at: str = field(default_factory=utc_now)
    updated_at: str = field(default_factory=utc_now)
    status: str = "active"
    writing_rules: list[dict[str, Any]] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class MaterialRecord:
    material_id: str
    matter_id: str
    title: str
    filename: str
    version: int
    content_hash: str
    stored_path: str
    text_path: str
    added_at: str
    tags: list[str] = field(default_factory=list)
    document_type: str = "material"
    supersedes: str | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class ChunkRecord:
    chunk_id: str
    owner_type: str
    owner_id: str
    matter_id: str
    text: str
    source_label: str
    created_at: str
    tags: list[str] = field(default_factory=list)
    vector: dict[str, float] = field(default_factory=dict)
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class KnowledgeRecord:
    knowledge_id: str
    matter_id: str
    title: str
    source: str
    source_type: str
    text_path: str
    added_at: str
    tags: list[str] = field(default_factory=list)
    validity_status: str = "active"
    effective_from: str | None = None
    effective_until: str | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


class CaseStore:
    """Filesystem-backed isolated case workspace."""

    def __init__(self, root: str | Path | None = None):
        env_root = os.environ.get("LEGALWORK_CASE_STORE")
        self.root = Path(root or env_root or (ROOT_DIR / ".legalwork_cases")).resolve()
        self.root.mkdir(parents=True, exist_ok=True)

    def matters_index_path(self) -> Path:
        return self.root / "matters.json"

    def list_matters(self) -> list[dict[str, Any]]:
        return load_json(self.matters_index_path(), [])

    def create_matter(
        self,
        name: str,
        team_id: str = "default",
        source_channel: str = "local",
        metadata: dict[str, Any] | None = None,
    ) -> Matter:
        matter_id = stable_id("matter")
        matter = Matter(
            matter_id=matter_id,
            name=name.strip() or matter_id,
            team_id=team_id or "default",
            source_channel=source_channel or "local",
            metadata=metadata or {},
        )
        self.matter_dir(matter_id).mkdir(parents=True, exist_ok=True)
        for sub in ("materials/files", "materials/text", "knowledge/text", "redacted", "drafts", "tasks"):
            (self.matter_dir(matter_id) / sub).mkdir(parents=True, exist_ok=True)
        self._save_matter(matter)
        matters = self.list_matters()
        matters.append(asdict(matter))
        write_json(self.matters_index_path(), matters)
        return matter

    def matter_dir(self, matter_id: str) -> Path:
        if not re.fullmatch(r"[A-Za-z0-9_-]+", matter_id or ""):
            raise ValueError("invalid matter_id")
        return self.root / matter_id

    def get_matter(self, matter_id: str) -> Matter:
        data = load_json(self.matter_dir(matter_id) / "matter.json", None)
        if not data:
            raise FileNotFoundError(f"matter not found: {matter_id}")
        return Matter(**data)

    def update_matter(self, matter: Matter) -> None:
        matter.updated_at = utc_now()
        self._save_matter(matter)
        matters = self.list_matters()
        for idx, item in enumerate(matters):
            if item.get("matter_id") == matter.matter_id:
                matters[idx] = asdict(matter)
                break
        write_json(self.matters_index_path(), matters)

    def _save_matter(self, matter: Matter) -> None:
        write_json(self.matter_dir(matter.matter_id) / "matter.json", asdict(matter))

    def materials_path(self, matter_id: str) -> Path:
        return self.matter_dir(matter_id) / "materials" / "index.json"

    def chunks_path(self, matter_id: str) -> Path:
        return self.matter_dir(matter_id) / "chunks.json"

    def knowledge_path(self, matter_id: str) -> Path:
        return self.matter_dir(matter_id) / "knowledge" / "index.json"

    def global_knowledge_dir(self) -> Path:
        path = self.root / "global_knowledge"
        (path / "text").mkdir(parents=True, exist_ok=True)
        return path

    def global_knowledge_path(self) -> Path:
        return self.global_knowledge_dir() / "index.json"

    def global_chunks_path(self) -> Path:
        return self.global_knowledge_dir() / "chunks.json"

    def list_materials(self, matter_id: str) -> list[dict[str, Any]]:
        return load_json(self.materials_path(matter_id), [])

    def list_chunks(self, matter_id: str) -> list[dict[str, Any]]:
        return load_json(self.chunks_path(matter_id), [])

    def list_knowledge(self, matter_id: str) -> list[dict[str, Any]]:
        return load_json(self.knowledge_path(matter_id), [])

    def list_global_knowledge(self) -> list[dict[str, Any]]:
        return load_json(self.global_knowledge_path(), [])

    def list_global_chunks(self) -> list[dict[str, Any]]:
        return load_json(self.global_chunks_path(), [])

    def _save_materials(self, matter_id: str, records: list[dict[str, Any]]) -> None:
        write_json(self.materials_path(matter_id), records)

    def _save_chunks(self, matter_id: str, chunks: list[dict[str, Any]]) -> None:
        write_json(self.chunks_path(matter_id), chunks)

    def _save_knowledge(self, matter_id: str, records: list[dict[str, Any]]) -> None:
        write_json(self.knowledge_path(matter_id), records)

    def _save_global_knowledge(self, records: list[dict[str, Any]]) -> None:
        write_json(self.global_knowledge_path(), records)

    def _save_global_chunks(self, chunks: list[dict[str, Any]]) -> None:
        write_json(self.global_chunks_path(), chunks)


class MaterialIndexer:
    def __init__(self, store: CaseStore):
        self.store = store

    def add_material(
        self,
        matter_id: str,
        file_path: str | Path | None = None,
        text: str | None = None,
        title: str = "",
        tags: list[str] | None = None,
        document_type: str = "material",
        metadata: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        self.store.get_matter(matter_id)
        matter_dir = self.store.matter_dir(matter_id)
        material_id = stable_id("mat")
        source_path = Path(file_path) if file_path else None
        extracted_text = text if text is not None else extract_text_from_path(source_path) if source_path else ""
        if not extracted_text.strip():
            raise ValueError("material text is empty")

        filename = source_path.name if source_path else f"{slugify_name(title, 'material')}.txt"
        safe_name = f"{material_id}_{slugify_name(filename, 'material.txt')}"
        stored_path = matter_dir / "materials" / "files" / safe_name
        if source_path and source_path.exists():
            shutil.copy2(source_path, stored_path)
            content_hash = sha256_file(source_path)
        else:
            stored_path.write_text(extracted_text, encoding="utf-8")
            content_hash = sha256_text(extracted_text)

        text_path = matter_dir / "materials" / "text" / f"{material_id}.txt"
        text_path.write_text(extracted_text, encoding="utf-8")

        records = self.store.list_materials(matter_id)
        logical_title = title.strip() or Path(filename).stem
        previous = [
            r
            for r in records
            if r.get("title") == logical_title or Path(r.get("filename", "")).stem == Path(filename).stem
        ]
        version = max([r.get("version", 0) for r in previous] or [0]) + 1
        supersedes = max(previous, key=lambda r: r.get("version", 0)).get("material_id") if previous else None
        record = MaterialRecord(
            material_id=material_id,
            matter_id=matter_id,
            title=logical_title,
            filename=filename,
            version=version,
            content_hash=content_hash,
            stored_path=str(stored_path),
            text_path=str(text_path),
            added_at=utc_now(),
            tags=tags or [],
            document_type=document_type or "material",
            supersedes=supersedes,
            metadata=metadata or {},
        )
        records.append(asdict(record))
        self.store._save_materials(matter_id, records)
        self._index_text(matter_id, "material", material_id, logical_title, extracted_text, tags or [], metadata or {})
        return asdict(record)

    def latest_materials(self, matter_id: str) -> list[dict[str, Any]]:
        latest: dict[str, dict[str, Any]] = {}
        for record in self.store.list_materials(matter_id):
            key = record.get("title") or Path(record.get("filename", "")).stem
            if key not in latest or record.get("version", 0) > latest[key].get("version", 0):
                latest[key] = record
        return list(latest.values())

    def search(self, matter_id: str, query: str, top_k: int = 8, latest_only: bool = True) -> list[dict[str, Any]]:
        query_vec = vectorize(query)
        latest_ids = {r["material_id"] for r in self.latest_materials(matter_id)}
        results = []
        for chunk in self.store.list_chunks(matter_id):
            if chunk.get("owner_type") != "material":
                continue
            if latest_only and chunk.get("owner_id") not in latest_ids:
                continue
            score = cosine(query_vec, chunk.get("vector", {}))
            if score > 0:
                item = dict(chunk)
                item["score"] = round(score, 6)
                results.append(item)
        results.sort(key=lambda item: item["score"], reverse=True)
        return results[:top_k]

    def _index_text(
        self,
        matter_id: str,
        owner_type: str,
        owner_id: str,
        source_label: str,
        text: str,
        tags: list[str],
        metadata: dict[str, Any],
    ) -> None:
        chunks = self.store.list_chunks(matter_id)
        chunks = [c for c in chunks if not (c.get("owner_type") == owner_type and c.get("owner_id") == owner_id)]
        for idx, chunk_text in enumerate(split_chunks(text)):
            chunk = ChunkRecord(
                chunk_id=f"{owner_id}_c{idx:04d}",
                owner_type=owner_type,
                owner_id=owner_id,
                matter_id=matter_id,
                text=chunk_text,
                source_label=source_label,
                created_at=utc_now(),
                tags=tags,
                vector=vectorize(chunk_text),
                metadata=metadata,
            )
            chunks.append(asdict(chunk))
        self.store._save_chunks(matter_id, chunks)


def split_chunks(text: str, max_chars: int = 1200) -> list[str]:
    if SemanticChunker:
        try:
            semantic = [c.text.strip() for c in SemanticChunker().chunk(text) if c.text.strip()]
            if semantic:
                out: list[str] = []
                for block in semantic:
                    out.extend(_split_long_block(block, max_chars))
                return out
        except Exception:
            pass
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    if not paragraphs:
        paragraphs = [text.strip()] if text.strip() else []
    out = []
    buf = ""
    for paragraph in paragraphs:
        if len(buf) + len(paragraph) + 2 <= max_chars:
            buf = f"{buf}\n\n{paragraph}".strip()
        else:
            if buf:
                out.append(buf)
            out.extend(_split_long_block(paragraph, max_chars))
            buf = ""
    if buf:
        out.append(buf)
    return out


def _split_long_block(text: str, max_chars: int) -> list[str]:
    if len(text) <= max_chars:
        return [text]
    return [text[i : i + max_chars] for i in range(0, len(text), max_chars)]


def build_chunk_records(
    matter_id: str,
    owner_type: str,
    owner_id: str,
    source_label: str,
    text: str,
    tags: list[str],
    metadata: dict[str, Any],
) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    for idx, chunk_text in enumerate(split_chunks(text)):
        chunk = ChunkRecord(
            chunk_id=f"{owner_id}_c{idx:04d}",
            owner_type=owner_type,
            owner_id=owner_id,
            matter_id=matter_id,
            text=chunk_text,
            source_label=source_label,
            created_at=utc_now(),
            tags=tags,
            vector=vectorize(chunk_text),
            metadata=metadata,
        )
        records.append(asdict(chunk))
    return records


def iter_knowledge_files(roots: list[Path], max_files: int = 800):
    count = 0
    for root in roots:
        if not root.exists():
            continue
        if root.is_file():
            candidates = [root]
        else:
            candidates = sorted(p for p in root.rglob("*") if p.is_file())
        for path in candidates:
            if path.name.startswith(".") or path.suffix.lower() not in KNOWLEDGE_FILE_EXTENSIONS:
                continue
            if any(part in {"node_modules", "__pycache__", ".git"} for part in path.parts):
                continue
            yield path
            count += 1
            if count >= max_files:
                return


def infer_source_type(path: Path) -> str:
    path_text = str(path)
    if "regulations" in path_text or "法规" in path_text or "法律" in path_text:
        return "regulation"
    if "template" in path_text or "模板" in path_text:
        return "template"
    if "case" in path_text or "案例" in path_text:
        return "case"
    if path.suffix.lower() == ".json":
        return "structured_reference"
    return "team_knowledge_file"


def relative_to_any(path: Path, roots: list[Path]) -> str:
    for root in roots:
        try:
            return str(path.resolve().relative_to(root.resolve()))
        except ValueError:
            continue
    return path.name


def dedupe_chunks(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    seen = set()
    out = []
    for item in sorted(items, key=lambda x: x.get("score", 0), reverse=True):
        key = (item.get("owner_type"), item.get("owner_id"), item.get("chunk_id"))
        if key in seen:
            continue
        seen.add(key)
        out.append(item)
    return out


def model_source_item(item: dict[str, Any]) -> dict[str, Any]:
    metadata = item.get("metadata") or {}
    return {
        "source_id": item.get("owner_id"),
        "chunk_id": item.get("chunk_id"),
        "scope": "global" if item.get("owner_type") == "global_knowledge" else item.get("matter_id"),
        "source_label": item.get("source_label"),
        "source_type": metadata.get("source_type") or item.get("owner_type"),
        "source_path": metadata.get("source_path") or metadata.get("source"),
        "relative_path": metadata.get("relative_path"),
        "validity_status": metadata.get("validity_status", "active"),
        "score": item.get("score", 0),
        "quote": item.get("text", ""),
    }


class KnowledgeBase:
    def __init__(self, store: CaseStore, knowledge_roots: list[str | Path] | None = None):
        self.store = store
        self.indexer = MaterialIndexer(store)
        self.knowledge_roots = [Path(p) for p in (default_knowledge_roots() if knowledge_roots is None else knowledge_roots)]

    def add_knowledge(
        self,
        matter_id: str,
        title: str,
        text: str,
        source: str = "manual",
        source_type: str = "team_experience",
        tags: list[str] | None = None,
        validity_status: str | None = None,
        effective_from: str | None = None,
        effective_until: str | None = None,
        metadata: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        self.store.get_matter(matter_id)
        status = validity_status or infer_validity_status(title, text, effective_until)
        knowledge_id = stable_id("kn")
        text_path = self.store.matter_dir(matter_id) / "knowledge" / "text" / f"{knowledge_id}.txt"
        text_path.write_text(text, encoding="utf-8")
        record = KnowledgeRecord(
            knowledge_id=knowledge_id,
            matter_id=matter_id,
            title=title.strip() or knowledge_id,
            source=source,
            source_type=source_type,
            text_path=str(text_path),
            added_at=utc_now(),
            tags=tags or infer_tags(title + "\n" + text),
            validity_status=status,
            effective_from=effective_from,
            effective_until=effective_until,
            metadata=metadata or {},
        )
        records = self.store.list_knowledge(matter_id)
        records.append(asdict(record))
        self.store._save_knowledge(matter_id, records)
        self.indexer._index_text(
            matter_id,
            "knowledge",
            knowledge_id,
            record.title,
            text,
            record.tags,
            {"source": source, "source_type": source_type, "validity_status": status},
        )
        return asdict(record)

    def import_file(
        self,
        matter_id: str,
        file_path: str | Path,
        title: str = "",
        source_type: str = "team_experience",
        tags: list[str] | None = None,
        metadata: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        path = Path(file_path)
        text = extract_text_from_path(path)
        if not text.strip():
            raise ValueError("knowledge text is empty")
        return self.add_knowledge(
            matter_id=matter_id,
            title=title or path.stem,
            text=text,
            source=str(path),
            source_type=source_type,
            tags=tags,
            metadata=metadata,
        )

    def search(self, matter_id: str, query: str, top_k: int = 8, include_inactive: bool = False) -> list[dict[str, Any]]:
        active_ids = {
            r["knowledge_id"]
            for r in self.store.list_knowledge(matter_id)
            if include_inactive or r.get("validity_status") == "active"
        }
        query_vec = vectorize(query)
        results = []
        for chunk in self.store.list_chunks(matter_id):
            if chunk.get("owner_type") != "knowledge" or chunk.get("owner_id") not in active_ids:
                continue
            score = cosine(query_vec, chunk.get("vector", {}))
            if score > 0:
                item = dict(chunk)
                item["score"] = round(score, 6)
                results.append(item)
        results.sort(key=lambda item: item["score"], reverse=True)
        return results[:top_k]

    def sync_global_sources(
        self,
        roots: list[str | Path] | None = None,
        max_files: int = 800,
        force: bool = False,
    ) -> dict[str, Any]:
        """Index local knowledge-base files so the agent can retrieve them automatically."""
        roots_to_scan = [Path(p) for p in (roots or self.knowledge_roots)]
        existing = self.store.list_global_knowledge()
        by_source = {r.get("source"): r for r in existing if r.get("source")}
        records = [dict(r) for r in existing]
        chunks = self.store.list_global_chunks()
        seen_sources: set[str] = set()
        added = 0
        updated = 0
        skipped = 0

        for path in iter_knowledge_files(roots_to_scan, max_files=max_files):
            source = str(path.resolve())
            seen_sources.add(source)
            try:
                size = path.stat().st_size
            except OSError:
                skipped += 1
                continue
            if size > MAX_KNOWLEDGE_FILE_BYTES:
                skipped += 1
                continue
            content_hash = sha256_file(path)
            current = by_source.get(source)
            if current and current.get("metadata", {}).get("content_hash") == content_hash and not force:
                continue
            text = extract_text_from_path(path)
            if not text.strip():
                skipped += 1
                continue
            status = infer_validity_status(path.stem, text)
            if current:
                knowledge_id = current["knowledge_id"]
                text_path = Path(current["text_path"])
                updated += 1
            else:
                knowledge_id = "gkn_" + hashlib.sha1(source.encode("utf-8")).hexdigest()[:16]
                text_path = self.store.global_knowledge_dir() / "text" / f"{knowledge_id}.txt"
                added += 1

            text_path.parent.mkdir(parents=True, exist_ok=True)
            text_path.write_text(text, encoding="utf-8")
            record = KnowledgeRecord(
                knowledge_id=knowledge_id,
                matter_id=GLOBAL_MATTER_ID,
                title=path.stem,
                source=source,
                source_type=infer_source_type(path),
                text_path=str(text_path),
                added_at=current.get("added_at", utc_now()) if current else utc_now(),
                tags=infer_tags(str(path) + "\n" + text[:2000]),
                validity_status=status,
                metadata={
                    "content_hash": content_hash,
                    "source_path": source,
                    "relative_path": relative_to_any(path, roots_to_scan),
                    "file_size": size,
                    "synced_at": utc_now(),
                },
            )
            records = [r for r in records if r.get("knowledge_id") != knowledge_id]
            records.append(asdict(record))
            chunks = [c for c in chunks if not (c.get("owner_type") == "global_knowledge" and c.get("owner_id") == knowledge_id)]
            chunks.extend(
                build_chunk_records(
                    matter_id=GLOBAL_MATTER_ID,
                    owner_type="global_knowledge",
                    owner_id=knowledge_id,
                    source_label=record.title,
                    text=text,
                    tags=record.tags,
                    metadata={
                        "source": source,
                        "source_path": source,
                        "source_type": record.source_type,
                        "validity_status": status,
                        "relative_path": record.metadata["relative_path"],
                    },
                )
            )

        for record in records:
            if record.get("source") and record["source"] not in seen_sources and Path(record["source"]).exists() is False:
                record["validity_status"] = "inactive"
                record.setdefault("metadata", {})["missing_from_disk"] = True

        self.store._save_global_knowledge(records)
        self.store._save_global_chunks(chunks)
        return {
            "success": True,
            "roots": [str(p) for p in roots_to_scan if p.exists()],
            "total_records": len(records),
            "total_chunks": len(chunks),
            "added": added,
            "updated": updated,
            "skipped": skipped,
        }

    def search_global(
        self,
        query: str,
        top_k: int = 8,
        include_inactive: bool = False,
        auto_sync: bool = True,
    ) -> list[dict[str, Any]]:
        if auto_sync and not self.store.list_global_knowledge():
            self.sync_global_sources()
        active_ids = {
            r["knowledge_id"]
            for r in self.store.list_global_knowledge()
            if include_inactive or r.get("validity_status") == "active"
        }
        query_vec = vectorize(query)
        results = []
        for chunk in self.store.list_global_chunks():
            if chunk.get("owner_type") != "global_knowledge" or chunk.get("owner_id") not in active_ids:
                continue
            score = cosine(query_vec, chunk.get("vector", {}))
            if score > 0:
                item = dict(chunk)
                item["score"] = round(score, 6)
                item["knowledge_scope"] = "global"
                results.append(item)
        results.sort(key=lambda item: item["score"], reverse=True)
        return results[:top_k]

    def agent_sources_for_task(self, matter_id: str, task: str, top_k: int = 12) -> dict[str, Any]:
        matter_hits = self.search(matter_id, task, top_k=max(3, top_k // 2))
        global_hits = self.search_global(task, top_k=top_k)
        merged = dedupe_chunks(matter_hits + global_hits)
        return {
            "query": task,
            "items": merged[:top_k],
            "matter_knowledge_count": len(matter_hits),
            "global_knowledge_count": len(global_hits),
            "policy": "Agent selects relevant active knowledge chunks automatically by local vector similarity and keeps source metadata for verification.",
        }


def infer_validity_status(title: str, text: str, effective_until: str | None = None) -> str:
    haystack = f"{title}\n{text[:2000]}"
    if re.search(r"废止|失效|过期|旧版|已否定|不再适用", haystack):
        return "inactive"
    if effective_until:
        try:
            if date.fromisoformat(effective_until) < date.today():
                return "inactive"
        except ValueError:
            pass
    return "active"


def infer_tags(text: str) -> list[str]:
    mapping = {
        "合同": "合同",
        "侵权": "侵权",
        "公司": "公司",
        "股权": "股权",
        "数据": "数据合规",
        "个人信息": "个人信息",
        "劳动": "劳动",
        "执行": "执行",
        "证据": "证据",
        "管辖": "管辖",
    }
    return [tag for keyword, tag in mapping.items() if keyword in text]


class FeedbackMemory:
    def __init__(self, store: CaseStore):
        self.store = store

    def add_rule(self, matter_id: str, feedback: str, scope: str = "matter", author: str = "") -> dict[str, Any]:
        matter = self.store.get_matter(matter_id)
        rule = {
            "rule_id": stable_id("rule"),
            "scope": scope or "matter",
            "author": author,
            "text": feedback.strip(),
            "created_at": utc_now(),
            "tags": infer_tags(feedback),
        }
        if not rule["text"]:
            raise ValueError("feedback is empty")
        matter.writing_rules.append(rule)
        self.store.update_matter(matter)
        return rule


class TemplateSkillRouter:
    DOC_TYPES = {
        "起诉书": ("complaint", ["起诉", "起诉状", "起诉书"]),
        "答辩状": ("answer", ["答辩", "答辩状"]),
        "代理意见": ("agency_opinion", ["代理意见", "庭后", "庭审意见"]),
        "法律意见书": ("legal_opinion", ["法律意见", "法律意见书"]),
        "尽调报告": ("due_diligence", ["尽调", "尽职调查"]),
        "合规报告": ("compliance_report", ["合规报告", "合规审查", "数据合规"]),
        "证据目录": ("evidence_catalog", ["证据目录", "证据清单"]),
    }
    SKILL_HINTS = {
        "complaint": ["fact_extraction", "legal_element_extraction", "document_drafting"],
        "answer": ["dispute_issue_identification", "evidence_argument_chain", "document_drafting"],
        "agency_opinion": ["argument_chain_construction", "evidence_evaluation", "document_drafting"],
        "legal_opinion": ["legal_research", "legal_risk_assessment", "legal-memo-generator"],
        "due_diligence": ["due_diligence", "contract_risk_review", "legal_risk_assessment"],
        "compliance_report": ["compliance_review", "data-compliance-ai-rd", "presidio-data-compliance"],
        "evidence_catalog": ["evidence-catalog", "timeline_generation", "evidence_evaluation"],
    }

    def __init__(self, skills_dir: str | Path | None = None):
        self.skills_dir = Path(skills_dir or (ROOT_DIR / "skills"))

    def route(self, task: str, matter: Matter | None = None) -> dict[str, Any]:
        task_text = task or ""
        doc_type = "legal_opinion"
        display = "法律意见书"
        for name, (code, keywords) in self.DOC_TYPES.items():
            if any(keyword in task_text for keyword in keywords):
                doc_type = code
                display = name
                break
        skills = [s for s in self.SKILL_HINTS.get(doc_type, []) if (self.skills_dir / s).exists()]
        if not skills:
            skills = [s for s in ("legal_research", "document_drafting") if (self.skills_dir / s).exists()]
        return {
            "document_type": doc_type,
            "document_type_label": display,
            "skills": skills,
            "template": self.find_template(doc_type),
            "writing_rules": matter.writing_rules if matter else [],
        }

    def find_template(self, doc_type: str) -> dict[str, Any]:
        candidates = []
        for root in (ROOT_DIR / "templates", ROOT_DIR / "skills"):
            if not root.exists():
                continue
            for path in root.rglob("*"):
                if path.is_file() and doc_type in path.stem and path.suffix.lower() in {".docx", ".md", ".txt"}:
                    candidates.append(path)
        if candidates:
            path = sorted(candidates, key=lambda p: len(str(p)))[0]
            return {"available": True, "path": str(path), "name": path.name}
        return {"available": False, "path": "", "name": f"default_{doc_type}"}


class SourceVerifier:
    def __init__(self, store: CaseStore):
        self.store = store

    def verify(self, matter_id: str, claims: list[dict[str, Any]]) -> dict[str, Any]:
        source_texts = self._source_texts(matter_id)
        checked = []
        for claim in claims:
            quote = (claim.get("quote") or "").strip()
            source_id = claim.get("source_id") or claim.get("owner_id")
            text = source_texts.get(source_id, "")
            verified = bool(quote and text and quote in text)
            checked.append({
                **claim,
                "verified": verified,
                "reason": "exact_quote_match" if verified else "quote_not_found_in_declared_source",
            })
        return {
            "success": True,
            "all_verified": all(item["verified"] for item in checked) if checked else False,
            "items": checked,
            "policy": "Only exact quotes found in the declared local source are marked verified.",
        }

    def _source_texts(self, matter_id: str) -> dict[str, str]:
        out: dict[str, str] = {}
        for record in self.store.list_materials(matter_id):
            out[record["material_id"]] = Path(record["text_path"]).read_text(encoding="utf-8", errors="replace")
        for record in self.store.list_knowledge(matter_id):
            out[record["knowledge_id"]] = Path(record["text_path"]).read_text(encoding="utf-8", errors="replace")
        for record in self.store.list_global_knowledge():
            out[record["knowledge_id"]] = Path(record["text_path"]).read_text(encoding="utf-8", errors="replace")
        return out


class DocumentDrafter:
    def __init__(self, store: CaseStore):
        self.store = store

    def build_plan(self, matter_id: str, task: str, context: list[dict[str, Any]], target_chars: int = 20000) -> dict[str, Any]:
        sections = infer_sections(task)
        per_section = max(1200, target_chars // max(1, len(sections)))
        return {
            "target_chars": target_chars,
            "split_required": target_chars >= 12000,
            "sections": [
                {
                    "index": idx + 1,
                    "title": title,
                    "target_chars": per_section,
                    "context_chunk_ids": [c.get("chunk_id") for c in context[idx:: max(1, len(sections))]][:6],
                }
                for idx, title in enumerate(sections)
            ],
            "stitching_rules": [
                "各分段先独立成稿，再统一术语、主体代称、证据编号和脚注格式。",
                "引用材料必须携带 source_id 与原文 quote，进入来源核验后再定稿。",
                "最终输出经过脱敏还原后再生成 Word/Markdown 文件。",
            ],
        }

    def draft_markdown(self, matter_id: str, task: str, bundle: dict[str, Any]) -> dict[str, Any]:
        matter_dir = self.store.matter_dir(matter_id)
        route = bundle.get("route", {})
        lines = [
            f"# {route.get('document_type_label', '法律文书')}初稿",
            "",
            "## 一、任务",
            task.strip(),
            "",
            "## 二、已选材料",
        ]
        for item in bundle.get("materials", [])[:10]:
            lines.append(f"- {item.get('source_label')}（相关度 {item.get('score', 0)}）")
        lines.extend(["", "## 三、可核验依据"])
        for item in bundle.get("knowledge", [])[:10]:
            quote = item.get("text", "")[:160].replace("\n", " ")
            lines.append(f"- 来源：{item.get('source_label')}；摘录：{quote}")
        lines.extend(["", "## 四、正文结构", ""])
        for section in bundle.get("long_document_plan", {}).get("sections", []):
            lines.append(f"### {section['title']}")
            lines.append("待接入线上大模型后，根据本节材料包展开生成。")
            lines.append("")
        path = matter_dir / "drafts" / f"{stable_id('draft')}.md"
        path.write_text("\n".join(lines), encoding="utf-8")
        return {"path": str(path), "format": "markdown"}


def infer_sections(task: str) -> list[str]:
    if any(k in task for k in ("代理意见", "庭后")):
        return ["案件事实", "争议焦点", "证据体系", "法律适用", "对对方观点的回应", "结论与请求"]
    if "起诉" in task:
        return ["当事人信息", "诉讼请求", "事实与理由", "证据目录", "管辖与程序事项"]
    if "答辩" in task:
        return ["答辩请求", "事实回应", "法律抗辩", "证据反驳", "结论"]
    if "尽调" in task:
        return ["项目概况", "资料清单", "核心风险", "整改建议", "附件索引"]
    return ["问题概述", "事实基础", "法律依据", "分析论证", "风险提示", "结论建议"]


class CourtAIPredictor:
    NODES = [
        ("jurisdiction", "管辖/受理"),
        ("facts", "事实认定"),
        ("evidence", "证据采信"),
        ("law", "法律适用"),
        ("liability", "责任承担"),
        ("remedy", "裁判结果/下一步"),
    ]

    def predict(self, facts: str, claims: str = "", evidence: list[dict[str, Any]] | None = None) -> dict[str, Any]:
        text = f"{facts}\n{claims}"
        evidence = evidence or []
        predictions = []
        for code, label in self.NODES:
            score = self._score_node(code, text, evidence)
            predictions.append({
                "node": code,
                "label": label,
                "confidence": score,
                "likely_result": self._label_result(score),
                "rationale": self._rationale(code, score),
                "next_move": self._next_move(code, score),
            })
        return {
            "success": True,
            "model": "local_rule_based_court_node_predictor_v1",
            "predictions": predictions,
            "disclaimer": "该结果用于对标法院 AI 判断节点的本地推演，不替代律师判断。",
        }

    def _score_node(self, code: str, text: str, evidence: list[dict[str, Any]]) -> float:
        score = 0.5
        if code == "jurisdiction" and re.search(r"管辖|住所地|合同履行地|约定管辖|仲裁", text):
            score += 0.2
        if code == "facts" and re.search(r"时间|地点|经过|聊天记录|付款|交付|签署", text):
            score += 0.18
        if code == "evidence":
            score += min(0.3, len(evidence) * 0.05)
            if re.search(r"原件|公证|鉴定|发票|银行流水|合同", text):
                score += 0.12
        if code == "law" and re.search(r"民法典|公司法|司法解释|法律依据|条", text):
            score += 0.16
        if code == "liability" and re.search(r"违约|侵权|过错|因果关系|损失", text):
            score += 0.18
        if code == "remedy" and re.search(r"请求|赔偿|返还|解除|继续履行|利息", text):
            score += 0.18
        if re.search(r"不确定|缺失|待补|无法证明|矛盾", text):
            score -= 0.18
        return round(max(0.05, min(0.95, score)), 2)

    def _label_result(self, score: float) -> str:
        if score >= 0.72:
            return "较可能支持"
        if score >= 0.45:
            return "存在不确定性"
        return "较可能不支持或需补强"

    def _rationale(self, code: str, score: float) -> str:
        if score >= 0.72:
            return f"{code} 节点已有较充分文本信号和证据信号。"
        if score >= 0.45:
            return f"{code} 节点有基础信号，但仍需要补强原文依据。"
        return f"{code} 节点支撑不足，建议优先补充可核验材料。"

    def _next_move(self, code: str, score: float) -> str:
        if score >= 0.72:
            return "进入文书论证，保留原文引用。"
        return "补充材料、明确争点，并重新跑来源核验。"


class CaseAIOrchestrator:
    """Local AI command center: case isolation, redaction, retrieval, routing."""

    def __init__(self, store: CaseStore | None = None, knowledge_roots: list[str | Path] | None = None):
        self.store = store or CaseStore()
        self.materials = MaterialIndexer(self.store)
        self.knowledge = KnowledgeBase(self.store, knowledge_roots=knowledge_roots)
        self.memory = FeedbackMemory(self.store)
        self.router = TemplateSkillRouter()
        self.verifier = SourceVerifier(self.store)
        self.drafter = DocumentDrafter(self.store)
        self.redaction = RedactionPipeline()
        self.renderer = RedactionRenderer()

    def capabilities(self) -> dict[str, Any]:
        return {
            "implemented": [
                "local_case_workspace",
                "case_context_isolation",
                "material_latest_version_index",
                "team_memory_rules",
                "local_redaction_before_model",
                "subject_consistent_redaction_mapping",
                "automatic_restoration_after_model",
                "knowledge_chunk_tag_vector_index",
                "global_knowledge_file_auto_sync",
                "agent_auto_knowledge_source_selection",
                "inactive_knowledge_filtering",
                "automatic_material_knowledge_selection",
                "source_exact_quote_verification",
                "template_skill_routing",
                "long_document_split_and_stitch_plan",
                "court_ai_node_prediction",
            ],
            "intentionally_ignored_for_now": [
                "enterprise_wechat_group_bot",
                "feishu_group_bot",
                "group_file_receive_callback",
                "group_result_push_callback",
                "remote_server_deployment_operations",
            ],
        }

    def plan_task(self, matter_id: str, task: str, target_chars: int = 20000) -> dict[str, Any]:
        matter = self.store.get_matter(matter_id)
        material_hits = self.materials.search(matter_id, task, top_k=10, latest_only=True)
        knowledge_selection = self.knowledge.agent_sources_for_task(matter_id, task, top_k=12)
        knowledge_hits = knowledge_selection["items"]
        route = self.router.route(task, matter)
        context = material_hits + knowledge_hits
        return {
            "matter": asdict(matter),
            "task": task,
            "route": route,
            "materials": material_hits,
            "knowledge": knowledge_hits,
            "knowledge_selection": knowledge_selection,
            "long_document_plan": self.drafter.build_plan(matter_id, task, context, target_chars),
        }

    def prepare_model_payload(
        self,
        matter_id: str,
        task: str,
        policy_name: str = "internal_legal_analysis",
        target_chars: int = 20000,
    ) -> dict[str, Any]:
        bundle = self.plan_task(matter_id, task, target_chars)
        raw_context = json.dumps(
            {
                "task": task,
                "route": bundle["route"],
                "materials": [model_source_item(x) for x in bundle["materials"]],
                "knowledge": [model_source_item(x) for x in bundle["knowledge"]],
                "knowledge_selection": bundle.get("knowledge_selection", {}),
                "writing_rules": bundle["route"].get("writing_rules", []),
                "long_document_plan": bundle["long_document_plan"],
            },
            ensure_ascii=False,
            indent=2,
        )
        result = self.redaction.process_text(raw_context, policy_name=policy_name)
        matter_dir = self.store.matter_dir(matter_id)
        mapping_path = matter_dir / "redacted" / f"{stable_id('mapping')}.mapping.enc"
        self.renderer.generate_mapping_file(
            result.get("mappings", []),
            str(mapping_path),
            encrypt=False,
            cluster_table=result.get("cluster_table", []),
        )
        return {
            "success": True,
            "matter_id": matter_id,
            "policy_name": policy_name,
            "redacted_payload": result["redacted_text"],
            "mapping_path": str(mapping_path),
            "entity_count": result.get("entity_count", 0),
            "bundle": bundle,
        }

    def restore_model_output(self, mapping_path: str, output_text: str) -> dict[str, Any]:
        restorer = RedactionRestorer(mapping_path)
        return restorer.restore_text(output_text)

    def run_local_task(self, matter_id: str, task: str, target_chars: int = 20000) -> dict[str, Any]:
        payload = self.prepare_model_payload(matter_id, task, target_chars=target_chars)
        draft = self.drafter.draft_markdown(matter_id, task, payload["bundle"])
        task_record = {
            "task_id": stable_id("task"),
            "matter_id": matter_id,
            "task": task,
            "created_at": utc_now(),
            "status": "prepared",
            "model_payload_mapping": payload["mapping_path"],
            "draft": draft,
            "note": "Local backend prepared redacted model payload and deterministic draft skeleton.",
        }
        write_json(self.store.matter_dir(matter_id) / "tasks" / f"{task_record['task_id']}.json", task_record)
        return {"success": True, "task": task_record, "payload": payload}
