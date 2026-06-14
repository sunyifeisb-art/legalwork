#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
import subprocess
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from ocr_text import (  # noqa: E402
    IMAGE_EXTENSIONS,
    OcrUnavailable,
    extract_image_text,
    extract_pdf_ocr_text,
)

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover - fallback handled at runtime
    PdfReader = None


def read_text(file_path: str, text: str) -> str:
    if file_path:
        path = Path(file_path)
        suffix = path.suffix.lower()
        if suffix == '.pdf':
            return read_pdf_text(path)
        if suffix in IMAGE_EXTENSIONS:
            return read_image_text(path)
        if suffix in {'.doc', '.docx'}:
            return read_office_text(path)
        return path.read_text(encoding='utf-8')
    return text


def read_pdf_text(path: Path) -> str:
    text = ''
    if PdfReader is not None:
        try:
            with path.open('rb') as handle:
                reader = PdfReader(handle)
                text = '\f'.join((page.extract_text() or '') for page in reader.pages).strip()
        except Exception:
            text = ''

    if not text and shutil.which('pdftotext'):
        run = subprocess.run(
            ['pdftotext', '-layout', '-enc', 'UTF-8', str(path), '-'],
            capture_output=True,
            text=True,
            check=False,
        )
        if run.returncode == 0:
            text = run.stdout.strip()

    if not text:
        try:
            text = extract_pdf_ocr_text(path)
        except OcrUnavailable as exc:
            raise SystemExit(f'无法从 PDF 中提取文本，且无法执行 OCR：{exc}') from exc
        except Exception as exc:
            raise SystemExit(f'无法从 PDF 中提取文本，OCR 失败：{exc}') from exc
    return text


def read_image_text(path: Path) -> str:
    try:
        text = extract_image_text(path)
    except OcrUnavailable as exc:
        raise SystemExit(f'图片审查需要 OCR：{exc}') from exc
    except Exception as exc:
        raise SystemExit(f'图片 OCR 失败：{exc}') from exc
    if not text:
        raise SystemExit('图片 OCR 未识别到可审查文本')
    return text


def read_office_text(path: Path) -> str:
    if path.suffix.lower() == '.docx':
        try:
            document = Document(str(path))
            text = '\n\n'.join(iter_docx_text(document)).strip()
        except Exception as exc:
            text = ''
        if text:
            return text
        if path.suffix.lower() != '.doc':
            raise SystemExit('无法从 DOCX 中提取文本，请确认文件内容有效')

    if not shutil.which('textutil'):
        raise SystemExit('当前环境不支持 .doc 解析，请先将文件另存为 .docx 或 .pdf')

    try:
        run = subprocess.run(
            ['textutil', '-convert', 'txt', '-stdout', str(path)],
            capture_output=True,
            text=True,
            check=False,
        )
    except Exception as exc:
        raise SystemExit(f'无法调用文档解析工具: {exc}') from exc

    if run.returncode != 0:
        message = run.stderr.strip() or f'textutil failed with exit code {run.returncode}'
        raise SystemExit(message)
    text = run.stdout.strip()
    if not text:
        raise SystemExit('无法从文档中提取文本，请确认文件内容有效')
    return text


def iter_docx_blocks(parent) -> list[Paragraph | Table]:
    container = parent.element.body if hasattr(parent, 'element') and hasattr(parent.element, 'body') else parent._tc
    blocks: list[Paragraph | Table] = []
    for child in container.iterchildren():
        if isinstance(child, CT_P):
            blocks.append(Paragraph(child, parent))
        elif isinstance(child, CT_Tbl):
            blocks.append(Table(child, parent))
    return blocks


def iter_table_text(table: Table) -> list[str]:
    parts: list[str] = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_parts = list(iter_docx_text(cell))
            if cell_parts:
                cells.append(' / '.join(cell_parts))
        if cells:
            parts.append(' | '.join(cells))
    return parts


def iter_docx_text(container) -> list[str]:
    parts: list[str] = []
    for block in iter_docx_blocks(container):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                parts.append(text)
        elif isinstance(block, Table):
            parts.extend(iter_table_text(block))

    if hasattr(container, 'sections'):
        for section in container.sections:
            for area in (section.header, section.footer):
                for paragraph in area.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        parts.append(text)
                for table in area.tables:
                    parts.extend(iter_table_text(table))

    return parts


def normalize(raw: str) -> str:
    text = raw.replace('\r\n', '\n').replace('\r', '\n')
    text = text.replace('\f', '\n')
    text = re.sub(r'[\t\u3000]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r' {2,}', ' ', text)
    return text.strip()


def segment(text: str, max_chars: int = 500) -> list[str]:
    paras = [p.strip() for p in text.split('\n\n') if p.strip()]
    out: list[str] = []
    buf = ''
    for p in paras:
        if not buf:
            buf = p
        elif len(buf) + 2 + len(p) <= max_chars:
            buf += '\n\n' + p
        else:
            out.append(buf)
            buf = p
    if buf:
        out.append(buf)
    if not out and text:
        out = [text[:max_chars]]
    return out


def extract_page_texts(file_path: str, raw: str) -> list[str]:
    path = Path(file_path) if file_path else None
    if path and path.suffix.lower() == '.pdf':
        raw_pages = [page.strip() for page in raw.split('\f')]
        return [normalize(page) for page in raw_pages if normalize(page)]
    if path and path.suffix.lower() in IMAGE_EXTENSIONS:
        normalized = normalize(raw)
        return [normalized] if normalized else []
    normalized = normalize(raw)
    return [normalized] if normalized else []


def segment_with_context(page_texts: list[str], max_chars: int = 500) -> list[dict]:
    contexts: list[dict] = []
    for page_index, page_text in enumerate(page_texts, start=1):
        page_segments = segment(page_text, max_chars=max_chars)
        for idx, seg in enumerate(page_segments, start=1):
            first_line = seg.splitlines()[0].strip() if seg.strip() else ''
            contexts.append({
                'page': page_index,
                'segment_index': len(contexts) + 1,
                'segment_in_page': idx,
                'label': first_line[:24] if first_line else f'第{idx}段',
                'text': seg,
            })
    return contexts


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument('--file', default='')
    parser.add_argument('--text', default='')
    parser.add_argument('--output', required=True)
    parser.add_argument('--max-chars', type=int, default=500)
    args = parser.parse_args()

    raw = read_text(args.file, args.text)
    if not raw.strip():
        raise SystemExit('empty input')

    normalized = normalize(raw)
    chunks = segment(normalized, max_chars=args.max_chars)
    page_texts = extract_page_texts(args.file, raw)
    segment_contexts = segment_with_context(page_texts, max_chars=args.max_chars)
    result = {
        'raw_length': len(raw),
        'normalized_length': len(normalized),
        'segment_count': len(chunks),
        'normalized_text': normalized,
        'segments': chunks,
        'page_count': len(page_texts),
        'segment_contexts': segment_contexts,
    }
    Path(args.output).write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding='utf-8')
    print(args.output)
    return 0


if __name__ == '__main__':
    raise SystemExit(main())
