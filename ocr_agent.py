#!/usr/bin/env python3
"""
LegalWork OCR Agent - 独立 OCR Agent 能力入口

可作为独立脚本运行，也可被其他 Agent 系统集成。
自动检测文件类型，选择最优 OCR 引擎，返回结构化识别结果。

用法：
  python3 ocr_agent.py scan <file_path>              # 对单文件执行 OCR
  python3 ocr_agent.py batch <dir_path>               # 批量 OCR 目录下所有文件
  python3 ocr_agent.py pipeline <file_path>            # 完整文处流水线（intake→OCR→LDIR）
  python3 ocr_agent.py engines                         # 列出可用 OCR 引擎
  python3 ocr_agent.py auto <file_path>                # 自动判断是否需要 OCR，按需执行
"""

import sys
import json
import os
from pathlib import Path

# 确保当前包可导入
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))


def scan_file(file_path: str, profile: str = "fast_local_ocr") -> dict:
    """
    对单个文件执行 OCR 识别。

    Args:
        file_path: 文件路径（支持 PDF、PNG、JPG、BMP、TIFF 等）
        profile: OCR profile（fast_local_ocr / legal_pdf_parser / complex_document / browser_local）

    Returns:
        dict: 包含 text、engine、confidence、blocks 等信息的结构化结果
    """
    from document.ocr.router import OCRRouter
    from document.intake.router import DocumentIntakeRouter

    path = Path(file_path)
    if not path.exists():
        return {"success": False, "error": f"文件不存在: {file_path}"}

    # 自动判断是否需要 OCR
    intake = DocumentIntakeRouter()
    decision = intake.route(str(path))

    # DOCX 特殊处理：直接读取文本
    if path.suffix.lower() in {".docx", ".doc"}:
        try:
            from docx import Document
            doc = Document(str(path))
            text_content = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return {
                "success": True,
                "text": text_content,
                "file_path": file_path,
                "file_type": decision.detected_type.value,
                "engine": "python-docx",
                "confidence": 0.95,
                "ocr_performed": False,
            }
        except Exception as e:
            # 兜底：继续走 OCR
            pass

    if not decision.ocr_required:
        # 文件已有文本层，直接读取
        try:
            text_content = path.read_text(encoding="utf-8", errors="replace")
        except Exception:
            text_content = path.read_text(encoding="gbk", errors="replace")

        return {
            "success": True,
            "text": text_content,
            "file_path": file_path,
            "file_type": decision.detected_type.value,
            "engine": "direct_read",
            "confidence": 1.0,
            "ocr_performed": False,
        }

    # 执行 OCR
    ocr = OCRRouter()
    result = ocr.process(str(path), profile_name=profile)

    if not result.text:
        return {
            "success": False,
            "file_path": file_path,
            "error": "OCR 未识别到文本",
            "available_engines": ocr.get_available_engines(),
            "file_type": decision.detected_type.value,
        }

    return {
        "success": True,
        "text": result.text,
        "markdown": result.markdown,
        "file_path": file_path,
        "file_type": decision.detected_type.value,
        "engine": result.engine,
        "confidence": result.confidence,
        "processing_time_ms": result.processing_time_ms,
        "blocks_count": len(result.blocks),
        "pages_count": len(result.pages),
        "ocr_performed": True,
        "profile": profile,
    }


def batch_scan(directory: str, profile: str = "fast_local_ocr") -> list:
    """
    批量 OCR 目录下所有可识别的文件。

    支持的格式：pdf, png, jpg, jpeg, bmp, tiff, tif, webp
    """
    supported = {".pdf", ".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp"}
    dir_path = Path(directory)
    if not dir_path.exists():
        return [{"success": False, "error": f"目录不存在: {directory}"}]

    results = []
    for f in sorted(dir_path.iterdir()):
        if f.is_file() and f.suffix.lower() in supported:
            result = scan_file(str(f), profile)
            results.append(result)

    return results


def pipeline(file_path: str) -> dict:
    """
    完整文档处理流水线：intake → OCR → LDIR → 结构化输出。

    适合需要结构化文档表示（带分块、坐标、置信度）的场景。
    """
    path = Path(file_path)
    if not path.exists():
        return {"success": False, "error": f"文件不存在: {file_path}"}

    from document.pipeline import DocumentPipeline

    pipeline_runner = DocumentPipeline()
    result = pipeline_runner.process(
        str(path),
        matter_id="standalone",
        matter_dir=path.parent,
    )
    return result


def auto_ocr(file_path: str) -> dict:
    """
    自动判断文件是否需要 OCR，按需执行。

    - 有文本层的 PDF/文本文件 → 直接读取
    - 扫描 PDF/图片 → 自动 OCR
    - 已有文本层但文本提取失败 → 兜底 OCR
    """
    path = Path(file_path)
    if not path.exists():
        return {"success": False, "error": f"文件不存在: {file_path}"}

    from document.intake.router import DocumentIntakeRouter
    from document.ocr.router import OCRRouter

    intake = DocumentIntakeRouter()
    decision = intake.route(str(path))

    if not decision.ocr_required:
        # 尝试直接读取
        for enc in ("utf-8", "utf-8-sig", "gbk", "gb18030"):
            try:
                text = path.read_text(encoding=enc)
                return {
                    "success": True,
                    "text": text,
                    "file_path": file_path,
                    "engine": f"direct_read({enc})",
                    "confidence": 1.0,
                    "ocr_performed": False,
                }
            except (UnicodeDecodeError, Exception):
                continue

        # 尝试读取 docx
        try:
            from docx import Document
            doc = Document(str(path))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return {
                "success": True,
                "text": text,
                "file_path": file_path,
                "engine": "docx",
                "confidence": 0.95,
                "ocr_performed": False,
            }
        except Exception:
            pass

        # 兜底：PyMuPDF
        try:
            import fitz
            doc = fitz.open(str(path))
            pages = [p.get_text("text") for p in doc]
            text = "\n\n".join(pages)
            if text.strip():
                return {
                    "success": True,
                    "text": text,
                    "file_path": file_path,
                    "engine": "pymupdf",
                    "confidence": 0.9,
                    "ocr_performed": False,
                }
        except Exception:
            pass

        # 兜底 OCR
        ocr = OCRRouter()
        result = ocr.process(str(path), profile_name="fast_local_ocr")
        if result.text:
            return {
                "success": True,
                "text": result.text,
                "file_path": file_path,
                "engine": f"ocr_fallback:{result.engine}",
                "confidence": result.confidence,
                "ocr_performed": True,
            }

        return {"success": False, "file_path": file_path, "error": "无法提取文本内容"}

    # 需要 OCR
    ocr = OCRRouter()
    result = ocr.process(str(path), profile_name="fast_local_ocr")
    if result.text:
        return {
            "success": True,
            "text": result.text,
            "file_path": file_path,
            "engine": result.engine,
            "confidence": result.confidence,
            "ocr_performed": True,
        }

    return {"success": False, "file_path": file_path, "error": "OCR 未识别到文本"}


def list_engines() -> dict:
    """列出当前环境可用的 OCR 引擎"""
    from document.ocr.router import OCRRouter
    ocr = OCRRouter()
    available = ocr.get_available_engines()
    return {
        "success": True,
        "engines": {
            "registered": list(ocr.adapters.keys()),
            "available": available,
        },
    }


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    command = sys.argv[1]

    if command == "scan":
        if len(sys.argv) < 3:
            print("用法: python3 ocr_agent.py scan <file_path> [profile]")
            sys.exit(1)
        profile = sys.argv[3] if len(sys.argv) > 3 else "fast_local_ocr"
        result = scan_file(sys.argv[2], profile)

    elif command == "batch":
        if len(sys.argv) < 3:
            print("用法: python3 ocr_agent.py batch <dir_path> [profile]")
            sys.exit(1)
        profile = sys.argv[3] if len(sys.argv) > 3 else "fast_local_ocr"
        results = batch_scan(sys.argv[2], profile)
        result = {"success": True, "count": len(results), "results": results}

    elif command == "pipeline":
        if len(sys.argv) < 3:
            print("用法: python3 ocr_agent.py pipeline <file_path>")
            sys.exit(1)
        result = pipeline(sys.argv[2])

    elif command == "auto":
        if len(sys.argv) < 3:
            print("用法: python3 ocr_agent.py auto <file_path>")
            sys.exit(1)
        result = auto_ocr(sys.argv[2])

    elif command == "engines":
        result = list_engines()

    else:
        print(f"未知命令: {command}")
        print(__doc__)
        sys.exit(1)

    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
